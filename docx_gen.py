import random
import time
from docx import Document, shared, enum
from utils import chunks, ensure_dir


batch_size = 10000
max_word_per_line = 12
min_word_per_line = 4
font_name = 'Times New Roman'
font_size = shared.Pt(13)


def line_processing(current_line, doc):
    list_word = current_line.split(' ')
    if len(list_word) > max_word_per_line:
        lines_list = list_word[0: max_word_per_line]
        p = doc.add_paragraph()
        list_index = range(len(lines_list))
        bold = random.choices(list_index, k=2)
        italic = random.choices(list_index, k=2)
        for idx, word in enumerate(lines_list):
            if idx in bold:
                p.add_run(word + " ").bold = True
            elif idx in italic:
                p.add_run(word + " ").italic = True
            else:
                p.add_run(word + " ")
        line_processing(' '.join(list_word[max_word_per_line + 1: len(list_word) - 1]), doc)
    else:
        if len(list_word) < min_word_per_line:
            return
        else:
            p = doc.add_paragraph()
            list_index = range(len(list_word))
            bold = random.choices(list_index, k=2)
            italic = random.choices(list_index, k=2)
            for idx, word in enumerate(list_word):
                if idx in bold:
                    p.add_run(word + " ").bold = True
                elif idx in italic:
                    p.add_run(word + " ").italic = True
                else:
                    p.add_run(word + " ")


def write_docx(lines, id_thread, outdir="out_gen_docx/"):
    start_time = time.time()
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = font_size
    prev_idx = 0
    ensure_dir(outdir)
    for idx, line in enumerate(lines):
        line = line.replace('\n', '')
        line_processing(line, document)
        if time.time() - start_time > 0.5:
            start_time = time.time()
            speed = (idx - prev_idx) * 2
            prev_idx = idx
            print("\rID_" + str(id_thread) + "_processed " + str(idx) 
                + " lines | speed: " + str(speed) + " item/s", end="")
        if (idx + 1) % batch_size == 0:
            document.save(outdir + 'processed_' + str(id_thread) + '_' + str(int(idx/batch_size)+1) + '.docx')
            document = Document()
            style = document.styles['Normal']
            font = style.font
            font.name = font_name
            font.size = font_size
    document.save(outdir + 'processed_' + str(id_thread) + '_' + str(int(idx/batch_size)+1) + '.docx')
